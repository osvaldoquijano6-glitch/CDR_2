"""Generador de protocolos (.docx) por cirugía sobre la plantilla del usuario.

Estrategia de fidelidad: el documento de salida se abre desde la plantilla
original (`normative/plantillas_usuario/Protocolo_*.docx`), de modo que la
portada, control de revisiones, definiciones, encabezados/pies, estilos
(Montserrat/Helvetica Neue, índigo #312E81) y capítulos finales quedan
idénticos. Solo se interviene:

1. Datos de portada y encabezado (proyecto, código, tipo, ubicación, fechas).
2. Columna "APLICA (SI/NO)" de la tabla maestra de pruebas.
3. El capítulo de pruebas, que se reconstruye desde el catálogo YAML
   (`normative/protocolos/catalogo_*.yaml`) insertando solo las secciones
   seleccionadas; las demás quedan como "(no aplica)" con su nota, igual que
   hace la plantilla. Las tablas de datos y los marcos de figura se clonan
   del XML original para conservar el formato exacto.

La numeración de pruebas es la del checklist / Anexo 5: 1–20 por unidad,
21–45 por central. `seccion` "2.1" pertenece a la prueba 2; en central,
"Prueba N -" del título equivale a la prueba N+20 del universo.
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from gcv.config.settings import NORMATIVE_DIR

PLANTILLAS = NORMATIVE_DIR / "plantillas_usuario"
CATALOGOS = NORMATIVE_DIR / "protocolos"

_CFG = {
    "central": {
        "plantilla": "Protocolo_Central_v1.2.1.docx",
        "catalogo": "catalogo_central.yaml",
        "inicio_h1": "3. Pruebas por Central Eléctrica",
        "fin_h1": "4. Datos de pruebas",
        "capitulo": "3",
        "offset_numero": 20,  # "Prueba 1" del protocolo = prueba 21 del universo
    },
    "unidad": {
        "plantilla": "Protocolo_Unidad_v1.2.1.docx",
        "catalogo": "catalogo_unidad.yaml",
        "inicio_h1": "4. Pruebas Por Unidad",
        "fin_h1": "5. Datos de pruebas",
        "capitulo": "4",
        "offset_numero": 0,
    },
}

_NO_APLICA_RE = re.compile(r"\s*[—\-(]*\s*no aplica\s*[)]?\s*$", re.IGNORECASE)


@dataclass
class ProyectoProtocolo:
    """Datos del proyecto que se estampan en la plantilla."""

    nombre_central: str                      # p.ej. "Central Eléctrica Solar Norte"
    codigo: str                              # clave corta para encabezados, p.ej. "SNO"
    proyecto: str = ""                       # p.ej. "GCE311_SNO_Solar"
    tipo: str = "B"                          # A/B/C/D
    tecnologia: str = "SINCRONA"             # SINCRONA / ASINCRONA
    ubicacion: str = ""
    fecha_pruebas: str = ""
    fecha_envio: str = ""
    # pruebas del universo 1–45 que aplican; el resto sale "(no aplica)"
    pruebas_aplican: set[int] = field(default_factory=set)
    # texto libre para pruebas con situación especial (numero → nota),
    # p.ej. 27: "No se cuenta con la infraestructura"
    notas: dict[int, str] = field(default_factory=dict)
    # sustituciones adicionales sobre los textos del catálogo
    reemplazos: dict[str, str] = field(default_factory=dict)

    def _reemplazos(self) -> dict[str, str]:
        base = {"CNK": self.codigo} if self.codigo else {}
        base.update(self.reemplazos)
        return base


def _cargar_catalogo(clave: str) -> dict:
    with (CATALOGOS / _CFG[clave]["catalogo"]).open(encoding="utf-8") as fh:
        return yaml.safe_load(fh)


def _numero_de(seccion: str, offset: int) -> int:
    return int(seccion.split(".")[0]) + offset


def _aplicar_reemplazos(texto: str, reemplazos: dict[str, str]) -> str:
    for viejo, nuevo in reemplazos.items():
        texto = texto.replace(viejo, nuevo)
    return texto


def _set_texto(p: Paragraph, texto: str) -> None:
    """Reescribe el texto conservando el formato del primer run."""
    if not p.runs:
        p.add_run(texto)
        return
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ""


def _reemplazar_en_parrafos(parrafos, reemplazos: dict[str, str]) -> None:
    for p in parrafos:
        if any(v in p.text for v in reemplazos):
            _set_texto(p, _aplicar_reemplazos(p.text, reemplazos))


def _portada(doc: Document, proyecto: ProyectoProtocolo) -> None:
    """Actualiza las líneas de portada por prefijo o valor conocido."""
    campos_prefijo = {
        "Proyecto:": proyecto.proyecto,
        "Fecha de las Pruebas:": proyecto.fecha_pruebas,
        "Fecha de Envío del Reporte:": proyecto.fecha_envio,
    }
    tecnologia = "SÍNCRONAS" if proyecto.tecnologia.upper().startswith("S") else "ASÍNCRONAS"
    for p in doc.paragraphs[:40]:
        t = p.text.strip()
        if not t:
            continue
        for prefijo, valor in campos_prefijo.items():
            if t.startswith(prefijo) and valor:
                _set_texto(p, f"{prefijo} {valor}")
        if t.startswith("PROTOCOLO DE PRUEBAS PARA PUESTA EN SERVICIO"):
            _set_texto(p, "PROTOCOLO DE PRUEBAS PARA PUESTA EN SERVICIO DE "
                          f"CENTRALES ELÉCTRICAS {tecnologia}")
        elif re.fullmatch(r"Central Tipo [A-D]", t):
            _set_texto(p, f"Central Tipo {proyecto.tipo}")
        elif t == "CATERPILLAR" and proyecto.nombre_central:
            _set_texto(p, proyecto.nombre_central)
        elif t == "Piedras Negras, Coahuila, México" and proyecto.ubicacion:
            _set_texto(p, proyecto.ubicacion)


def _encabezados_pies(doc: Document, reemplazos: dict[str, str]) -> None:
    for sec in doc.sections:
        for parte in (sec.header, sec.footer, sec.first_page_header,
                      sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            if parte is not None:
                _reemplazar_en_parrafos(parte.paragraphs, reemplazos)


def _tabla_maestra(doc: Document, proyecto: ProyectoProtocolo) -> None:
    """Rellena la columna APLICA (SI/NO) según la selección del proyecto."""
    for t in doc.tables:
        if "APLICA (SI/NO)" not in t.rows[0].cells[-1].text:
            continue
        for fila in t.rows[2:]:
            celdas = fila.cells
            num_txt = celdas[0].text.strip()
            if not num_txt.isdigit():
                continue  # banda de categoría
            numero = int(num_txt)
            nota = proyecto.notas.get(numero)
            valor = nota if nota else ("SI" if numero in proyecto.pruebas_aplican else "No Aplica")
            celda = celdas[-1]
            _set_texto(celda.paragraphs[0], valor)
            for extra in celda.paragraphs[1:]:
                _set_texto(extra, "")
        return


def _firma_tabla(t: Table) -> str:
    return " ".join(t.rows[0].cells[0].text.split())[:60]


def _clonar_tablas_capitulo(doc: Document, cfg: dict) -> dict[str, object]:
    """Copia el XML de las tablas del capítulo, indexadas por su primera celda."""
    clones: dict[str, object] = {}
    dentro = False
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.style.name == "Heading 1":
                if p.text.strip() == cfg["inicio_h1"]:
                    dentro = True
                elif dentro:
                    break
        elif child.tag == qn("w:tbl") and dentro:
            t = Table(child, doc)
            clones.setdefault(_firma_tabla(t), deepcopy(child))
            if "LA PRUEBA APLICA" in t.rows[0].cells[0].text:
                clones.setdefault("__aplicabilidad__", deepcopy(child))
            if len(t.rows) == 1 and len(t.columns) == 1:
                clones.setdefault("__placeholder__", deepcopy(child))
    return clones


def _vaciar_capitulo(doc: Document, cfg: dict):
    """Elimina desde la primera sección de prueba hasta antes del H1 final.

    Devuelve el elemento ancla (el H1 final) ante el cual insertar.
    """
    dentro = False
    borrar = []
    ancla = None
    en_pruebas = False
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            estilo, texto = p.style.name, p.text.strip()
            if estilo == "Heading 1":
                if texto == cfg["inicio_h1"]:
                    dentro = True
                    continue
                if dentro:
                    ancla = child
                    break
                continue
            if dentro and estilo == "Heading 2":
                en_pruebas = True
        if dentro and en_pruebas:
            borrar.append(child)
    for el in borrar:
        el.getparent().remove(el)
    return ancla


def _marcar_aplicabilidad(tbl_el, tipo: str, aplica: bool) -> None:
    """Pone la X bajo el tipo de central del proyecto (fila de marcas)."""
    doc_tbl = tbl_el
    filas = doc_tbl.findall(qn("w:tr"))
    if len(filas) < 3 or not aplica:
        return
    columnas = {"A": 0, "B": 2, "C": 4, "D": 6}
    idx = columnas.get(tipo.upper())
    if idx is None:
        return
    celdas = filas[2].findall(qn("w:tc"))
    if idx < len(celdas):
        # escribe la X en el primer párrafo de la celda
        p = celdas[idx].find(qn("w:p"))
        if p is not None:
            r = p.makeelement(qn("w:r"), {})
            t = p.makeelement(qn("w:t"), {})
            t.text = "X"
            r.append(t)
            p.append(r)


class _Insertador:
    """Inserta párrafos/tablas justo antes del ancla, conservando el orden."""

    def __init__(self, doc: Document, ancla):
        self.doc, self.ancla = doc, ancla

    def parrafo(self, texto: str, estilo: str) -> Paragraph:
        try:
            p = self.doc.add_paragraph(texto, style=estilo)
        except KeyError:
            p = self.doc.add_paragraph(texto)
        self.ancla.addprevious(p._p)
        return p

    def tabla_clonada(self, tbl_el) -> object:
        nuevo = deepcopy(tbl_el)
        self.ancla.addprevious(nuevo)
        return nuevo

    def tabla_datos(self, filas: list[list[str]]) -> None:
        ncols = max(len(f) for f in filas)
        t = self.doc.add_table(rows=len(filas), cols=ncols)
        t.style = "Table Grid"
        for i, fila in enumerate(filas):
            for j, valor in enumerate(fila):
                t.rows[i].cells[j].text = valor or ""
        self.ancla.addprevious(t._tbl)


def _texto_tabla_clonada(tbl_el, reemplazos: dict[str, str]) -> None:
    for t_el in tbl_el.iter(qn("w:t")):
        if t_el.text:
            t_el.text = _aplicar_reemplazos(t_el.text, reemplazos)


def _insertar_prueba(ins: _Insertador, entrada: dict, cfg: dict,
                     proyecto: ProyectoProtocolo, clones: dict) -> None:
    numero = _numero_de(entrada["seccion"], cfg["offset_numero"])
    aplica = numero in proyecto.pruebas_aplican
    reemplazos = proyecto._reemplazos()

    titulo_base = _NO_APLICA_RE.sub("", entrada["titulo"]).strip()
    sufijo = "" if aplica else " (no aplica)"
    ins.parrafo(f"{cfg['capitulo']}.{entrada['seccion']}. {titulo_base}{sufijo}",
                "Heading 2")

    if not aplica:
        tbl = clones.get("__aplicabilidad__")
        if tbl is not None:
            ins.tabla_clonada(tbl)
        nota = proyecto.notas.get(numero, "")
        detalle = f" {nota.rstrip('.')}." if nota else ""
        ins.parrafo(f"No aplica para la {proyecto.nombre_central}.{detalle}", "Normal")
        return

    for item in entrada.get("flujo", []):
        if "tabla" in item and item["tabla"] == "aplicabilidad":
            tbl = clones.get("__aplicabilidad__")
            if tbl is not None:
                nuevo = ins.tabla_clonada(tbl)
                _marcar_aplicabilidad(nuevo, proyecto.tipo, aplica=True)
        elif "placeholder" in item:
            texto = _aplicar_reemplazos(item["placeholder"], reemplazos)
            tbl = clones.get(_firma_tabla_texto(item["placeholder"]))
            if tbl is None:
                tbl = clones.get("__placeholder__")
            if tbl is not None:
                nuevo = ins.tabla_clonada(tbl)
                _texto_tabla_clonada(nuevo, reemplazos)
            else:
                ins.parrafo(texto, "Normal")
        elif "tabla_datos" in item:
            firma = _firma_tabla_texto(item["tabla_datos"][0][0] if item["tabla_datos"] else "")
            tbl = clones.get(firma)
            if tbl is not None:
                nuevo = ins.tabla_clonada(tbl)
                _texto_tabla_clonada(nuevo, reemplazos)
            else:
                ins.tabla_datos(item["tabla_datos"])
        elif "estilo" in item:
            texto = _aplicar_reemplazos(item.get("texto", ""), reemplazos)
            if item.get("imagen_incrustada") and not texto:
                tbl = clones.get("__placeholder__")
                if tbl is not None:
                    nuevo = ins.tabla_clonada(tbl)
                    for t_el in nuevo.iter(qn("w:t")):
                        t_el.text = ""
                    p = nuevo.find(".//" + qn("w:p"))
                    if p is not None:
                        r = p.makeelement(qn("w:r"), {})
                        te = p.makeelement(qn("w:t"), {})
                        te.text = "[Espacio reservado para evidencia del proyecto]"
                        r.append(te)
                        p.append(r)
                continue
            ins.parrafo(texto, item["estilo"])


def _firma_tabla_texto(texto: str) -> str:
    return " ".join((texto or "").split())[:60]


def generar_protocolo(clase: str, proyecto: ProyectoProtocolo, destino: Path) -> Path:
    """Genera el protocolo ('central' o 'unidad') fiel a la plantilla.

    `destino` puede ser carpeta o ruta .docx.
    """
    if clase not in _CFG:
        raise ValueError(f"clase debe ser 'central' o 'unidad', no {clase!r}")
    cfg = _CFG[clase]
    catalogo = _cargar_catalogo(clase)
    doc = Document(str(PLANTILLAS / cfg["plantilla"]))

    reemplazos = proyecto._reemplazos()
    _portada(doc, proyecto)
    _encabezados_pies(doc, reemplazos)
    _tabla_maestra(doc, proyecto)

    clones = _clonar_tablas_capitulo(doc, cfg)
    ancla = _vaciar_capitulo(doc, cfg)
    if ancla is None:
        raise RuntimeError(f"No se encontró el capítulo final {cfg['fin_h1']!r} en la plantilla")

    ins = _Insertador(doc, ancla)
    for entrada in catalogo["pruebas"]:
        if "grupo" in entrada:
            ins.parrafo(entrada["grupo"], "Heading 2")
        else:
            _insertar_prueba(ins, entrada, cfg, proyecto, clones)

    destino = Path(destino)
    if destino.suffix.lower() != ".docx":
        destino.mkdir(parents=True, exist_ok=True)
        destino = destino / f"Protocolo_{clase.capitalize()}_{proyecto.codigo or 'GCV'}.docx"
    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    return destino
