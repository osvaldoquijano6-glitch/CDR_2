"""Protocolo estilo v2 — rediseño regulatorio aprobado por el usuario.

REESTRUCTURA, NO ELIMINA: la estructura documental completa se extrae de la
plantilla original del usuario y se conserva íntegra — portada, control de
revisiones, tabla de contenido, definiciones y abreviaciones, objetivo,
alcance, capítulo de pruebas, datos de pruebas, certificados de calibración
y referencias. Solo cambia el tratamiento visual:

  * banda de marca azul eléctrico + eyebrow normativo en portada
  * placas de sección numeradas en azul
  * tablas con cabecera sólida azul y retícula recesiva
  * tabla maestra con chips SI / NO APLICA
  * tabla de contenido como campo real de Word (F9 la actualiza)

Paleta (Variante A aprobada): eléctrico #0EA5D8 · eléctrico oscuro #0B7FA8 ·
marino #1B3A6B · apagado #7FB8D9 · cuerpo #2C3A48.

Tipografías UNIVERSALES para que cualquier revisor las tenga instaladas:
Arial (display/encabezados) y Calibri (cuerpo). Nada de fuentes que el
usuario final deba instalar.

Regla dura: la columna APLICA solo lleva SI / NO / APLICA / NO APLICA. Las
notas libres provienen del usuario (visita en sitio, CENACE, dictamen); este
generador solo las estampa si vienen en `proyecto.notas`, nunca las redacta.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from gcv.protocolos.builder import (PLANTILLAS, ProyectoProtocolo,
                                    _aplicar_reemplazos, _cargar_catalogo,
                                    _numero_de, _CFG)
from gcv.protocolos.checklist import universo_pruebas

# ── paleta aprobada (Variante A) ────────────────────────────────────────────
ELECTRIC = "0EA5D8"
ELECTRIC_DARK = "0B7FA8"
NAVY = "1B3A6B"
MUTED = "7FB8D9"
BODY = "2C3A48"
LABEL = "7A7060"
CHIP_SI_BG = "E1F3FA"
CHIP_NO_BG = "EDEFF1"
CHIP_NO_FG = "7D8A94"
ROW_LINE = "E4EDF2"
RULE = "CFE4EF"

# tipografías universales (instaladas con Office/Windows/macOS)
DISPLAY = "Arial"
CUERPO = "Calibri"


# ── helpers XML ─────────────────────────────────────────────────────────────
def _rgb(hexval: str) -> RGBColor:
    return RGBColor.from_string(hexval)


def _cell_bg(cell, hexval: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexval)
    tc_pr.append(shd)


def _run_spacing(run, val: int = 30) -> None:
    r_pr = run._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(val))
    r_pr.append(sp)


def _run_bg(run, hexval: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexval)
    r_pr.append(shd)


def _p_bottom_border(p, hexval: str = RULE, sz: int = 6) -> None:
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hexval)
    borders.append(bottom)
    p_pr.append(borders)


def _cell_borders(cell, bottom: str | None = ROW_LINE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "right"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    el = OxmlElement("w:bottom")
    if bottom:
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), bottom)
    else:
        el.set(qn("w:val"), "nil")
    borders.append(el)
    tc_pr.append(borders)


def _row_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _texto(p, texto: str, *, font: str = CUERPO, size: float = 10.5,
           bold: bool = False, italic: bool = False, color: str = BODY,
           spacing: int | None = None, bg: str | None = None):
    run = p.add_run(texto)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = _rgb(color)
    if spacing:
        _run_spacing(run, spacing)
    if bg:
        _run_bg(run, bg)
    return run


# ── extracción de la estructura desde la plantilla del usuario ──────────────
def _extraer_plantilla(clase: str) -> list[tuple]:
    """Recorre el cuerpo de la plantilla original en orden y devuelve items:

    ("h1"|"h2", texto) · ("p"|"bullet", texto) · ("figura", pie_de_figura) ·
    ("tabla", filas). Los índices de la tabla de contenido original ("toc N")
    se omiten porque el v2 usa un campo TOC real.
    """
    doc = Document(str(PLANTILLAS / _CFG[clase]["plantilla"]))
    items: list[tuple] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            texto = p.text.strip()
            estilo = p.style.name or ""
            if estilo.lower().startswith("toc"):
                continue
            tiene_img = (child.find(".//" + qn("w:drawing")) is not None
                         or child.find(".//" + qn("w:pict")) is not None)
            if estilo == "Heading 1":
                items.append(("h1", texto))
            elif estilo == "Heading 2":
                items.append(("h2", texto))
            elif tiene_img:
                items.append(("figura", texto))
            elif texto:
                es_bullet = child.find(".//" + qn("w:numPr")) is not None
                items.append(("bullet" if es_bullet else "p", texto))
        elif child.tag == qn("w:tbl"):
            t = Table(child, doc)
            items.append(("tabla", [[c.text.strip() for c in r.cells]
                                    for r in t.rows]))
    return items


def _es_tabla_maestra(filas: list[list[str]]) -> bool:
    """Central: fila 0 termina en 'APLICA (SI/NO)'. Unidad: banda de título
    en fila 0 y encabezados ('Prueba…Aplica') en fila 1."""
    for fila in filas[:2]:
        if (fila and "APLICA" in (fila[-1] or "").upper()
                and "LA PRUEBA APLICA" not in (fila[0] or "").upper()):
            return True
    return False


def _es_tabla_aplicabilidad(filas: list[list[str]]) -> bool:
    return bool(filas) and "LA PRUEBA APLICA" in (filas[0][0] or "")


# ── bloques visuales v2 ─────────────────────────────────────────────────────
def _banda_marca(doc: Document) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _cell_bg(cell, ELECTRIC)
    _cell_borders(cell, bottom=None)
    cell.paragraphs[0].add_run(" ").font.size = Pt(2)
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), "140")
    h.set(qn("w:hRule"), "exact")
    tr_pr.append(h)


def _estilos_encabezados(doc: Document) -> None:
    """Ajusta los estilos Heading para que el campo TOC funcione y el
    documento sea reestilable desde Word (pestaña Estilos)."""
    for nombre, size, color in (("Heading 1", 13, NAVY),
                                ("Heading 2", 11.5, ELECTRIC_DARK),
                                ("Heading 3", 10.5, ELECTRIC_DARK)):
        st = doc.styles[nombre]
        st.font.name = DISPLAY
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = _rgb(color)
    normal = doc.styles["Normal"]
    normal.font.name = CUERPO
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(BODY)


_NUM_H1 = re.compile(r"^(\d+)\.?\s*(.+)$")


def _h1(doc: Document, texto: str, muted: bool = False) -> None:
    """Heading 1 real (aparece en la tabla de contenido) con placa azul."""
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    m = _NUM_H1.match(texto)
    if m:
        _texto(p, f" {m.group(1)} ", font=DISPLAY, size=11.5, bold=True,
               color="FFFFFF", bg=(MUTED if muted else ELECTRIC))
        _texto(p, f"  {m.group(2)}", font=DISPLAY, size=13, bold=True,
               color=(MUTED if muted else NAVY))
    else:
        _texto(p, texto, font=DISPLAY, size=13, bold=True,
               color=(MUTED if muted else NAVY))
    _p_bottom_border(p, RULE, sz=6)


def _h2(doc: Document, numero: str, titulo: str, muted: bool = False) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if numero:
        _texto(p, f" {numero} ", font=DISPLAY, size=9.5, bold=True,
               color="FFFFFF", bg=(MUTED if muted else ELECTRIC))
        _texto(p, f"  {titulo}", font=DISPLAY, size=11.5, bold=True,
               color=(MUTED if muted else NAVY))
    else:
        _texto(p, titulo, font=DISPLAY, size=11.5, bold=True,
               color=(MUTED if muted else NAVY))
    _p_bottom_border(p, RULE, sz=4)


def _parrafo(doc: Document, texto: str, bullet: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bullet:
        p.paragraph_format.left_indent = Cm(0.6)
        _texto(p, "•  ", size=10.5, bold=True, color=ELECTRIC_DARK)
    _texto(p, texto, size=10.5, color=BODY)


def _figura_placeholder(doc: Document, pie: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    _cell_bg(c, "F4FAFD")
    tc_pr = c._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "right", "bottom"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "dashed")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), MUTED)
        borders.append(el)
    tc_pr.append(borders)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _texto(p, f"[Espacio para {pie or 'figura'}]", size=8.5,
           color=CHIP_NO_FG, italic=True)
    if pie:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        _texto(cap, pie, size=8.5, color=LABEL, italic=True)


def _tabla_v2(doc: Document, filas: list[list[str]],
              primera_es_cabecera: bool = True) -> None:
    """Tabla genérica con el tratamiento v2 (cabecera azul, retícula suave).

    Colapsa celdas combinadas repetidas de la plantilla original.
    """
    limpias = []
    for fila in filas:
        vista, previa = [], None
        for v in fila:
            vista.append("" if v == previa else v)
            previa = v
        limpias.append(vista)
    ncols = max(len(f) for f in limpias)
    t = doc.add_table(rows=0, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, fila in enumerate(limpias):
        row = t.add_row()
        es_cab = primera_es_cabecera and i == 0
        if es_cab:
            _row_header_repeat(row)
        for j in range(ncols):
            c = row.cells[j]
            valor = fila[j] if j < len(fila) else ""
            if es_cab:
                _cell_bg(c, ELECTRIC)
                _cell_borders(c, bottom=None)
                _texto(c.paragraphs[0], valor, font=DISPLAY, size=8.5,
                       bold=True, color="FFFFFF")
            else:
                _cell_borders(c)
                _texto(c.paragraphs[0], valor, size=9, color=BODY)


def _chip(cell, texto: str, aplica: bool) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _texto(p, f" {texto} ", font=DISPLAY, size=8, bold=True,
           color=(ELECTRIC_DARK if aplica else CHIP_NO_FG),
           bg=(CHIP_SI_BG if aplica else CHIP_NO_BG))


def _tabla_maestra_v2(doc: Document, proyecto: ProyectoProtocolo,
                      clase: str) -> None:
    universo = universo_pruebas()
    if clase == "central":
        filas = [u for u in universo if u["numero"] >= 21]
    else:
        filas = [u for u in universo if u["numero"] <= 20]

    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    anchos = (Cm(1.1), Cm(5.2), Cm(8.4), Cm(2.6))
    hdr = t.rows[0]
    _row_header_repeat(hdr)
    for j, texto in enumerate(("No.", "Prueba", "Criterio de aceptación",
                               "Aplica")):
        c = hdr.cells[j]
        c.width = anchos[j]
        _cell_bg(c, ELECTRIC)
        _cell_borders(c, bottom=None)
        p = c.paragraphs[0]
        if j in (0, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _texto(p, texto, font=DISPLAY, size=9, bold=True, color="FFFFFF")

    for u in filas:
        numero = u["numero"]
        nota = proyecto.notas.get(numero, "").strip()
        aplica = numero in proyecto.pruebas_aplican
        estado = nota if nota else ("SI" if aplica else "NO APLICA")
        color_txt = BODY if aplica else MUTED
        row = t.add_row()
        for j, texto in enumerate((str(numero), u["nombre"], u["criterio"])):
            c = row.cells[j]
            c.width = anchos[j]
            _cell_borders(c)
            p = c.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _texto(p, texto, size=9.5, bold=True,
                       color=(ELECTRIC_DARK if aplica else MUTED))
            else:
                _texto(p, texto, size=9.5, color=color_txt)
        c = row.cells[3]
        c.width = anchos[3]
        _cell_borders(c)
        _chip(c, estado, aplica)


def _tabla_aplicabilidad_v2(doc: Document, tipo: str) -> None:
    t = doc.add_table(rows=2, cols=8)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    etiquetas = ("A", "", "B", "", "C", "", "D", "")
    for j in range(8):
        c0, c1 = t.rows[0].cells[j], t.rows[1].cells[j]
        for c in (c0, c1):
            _cell_borders(c)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if etiquetas[j]:
            _cell_bg(c0, CHIP_SI_BG)
            _texto(c0.paragraphs[0], f"Tipo {etiquetas[j]}", font=DISPLAY,
                   size=8.5, bold=True, color=ELECTRIC_DARK)
            marca = "X" if etiquetas[j] == tipo.upper() else ""
            _texto(c1.paragraphs[0], marca, size=10, bold=True, color=NAVY)


# ── portada ─────────────────────────────────────────────────────────────────
def _portada(doc: Document, proyecto: ProyectoProtocolo,
             items: list[tuple], alcance_pill: str) -> None:
    """Portada v2 con TODAS las líneas de la portada original, re-estilizadas.

    Las líneas se toman de la plantilla y se sustituyen los datos del
    proyecto con las mismas reglas que el generador fiel (builder._portada).
    """
    reemplazos = proyecto._reemplazos()
    tecnologia = ("SÍNCRONAS" if proyecto.tecnologia.upper().startswith("S")
                  else "ASÍNCRONAS")

    _banda_marca(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    _texto(p, "CÓDIGO DE RED 2.0 · RES/550/2021 · MANUAL INTE · POC ANEXO 5",
           font=DISPLAY, size=8, bold=True, color=ELECTRIC_DARK, spacing=40)

    for tipo_item, texto in items:
        if tipo_item != "p":
            continue
        t = texto.strip()
        if t.startswith("PROTOCOLO DE PRUEBAS"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            _texto(p, "PROTOCOLO DE PRUEBAS PARA PUESTA EN SERVICIO DE ",
                   font=DISPLAY, size=16, bold=True, color=NAVY)
            _texto(p, f"CENTRALES ELÉCTRICAS {tecnologia}",
                   font=DISPLAY, size=16, bold=True, color=ELECTRIC)
        elif re.fullmatch(r"Central Tipo [A-D]", t):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            _texto(p, f" CENTRAL TIPO {proyecto.tipo.upper()} ", font=DISPLAY,
                   size=9, bold=True, color="FFFFFF", bg=ELECTRIC)
            if alcance_pill:
                _texto(p, "  ")
                _texto(p, f" {alcance_pill.upper()} ", font=DISPLAY, size=9,
                       bold=True, color=ELECTRIC_DARK, bg=CHIP_SI_BG)
        elif t == "CATERPILLAR":
            p = doc.add_paragraph()
            _texto(p, proyecto.nombre_central or t, font=DISPLAY, size=20,
                   bold=True, color=NAVY)
        elif t == "Piedras Negras, Coahuila, México":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            _texto(p, proyecto.ubicacion or t, size=11, color=LABEL)
        elif t == "PRUEBAS DE CÓDIGO DE RED":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _texto(p, t, font=DISPLAY, size=10, bold=True,
                   color=ELECTRIC_DARK, spacing=30)
            _p_bottom_border(p, RULE, sz=8)
        elif t == "CONTROL DE REVISIONES":
            break
        elif ":" in t:
            etiqueta, _, valor = t.partition(":")
            valor = valor.strip()
            campos = {"Proyecto": proyecto.proyecto,
                      "Fecha de las Pruebas": proyecto.fecha_pruebas,
                      "Fecha de Envío del Reporte": proyecto.fecha_envio}
            valor = campos.get(etiqueta.strip()) or valor
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _texto(p, f"{etiqueta.strip().upper()}  ", font=DISPLAY, size=8,
                   bold=True, color=LABEL, spacing=30)
            _texto(p, _aplicar_reemplazos(valor, reemplazos), size=11,
                   bold=True, color=NAVY)
        else:
            p = doc.add_paragraph()
            _texto(p, _aplicar_reemplazos(t, reemplazos), size=11.5,
                   bold=True, color=BODY)
    doc.add_page_break()


def _tabla_contenido(doc: Document) -> None:
    _h1(doc, "Tabla de contenido")
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = ("Haga clic aquí y presione F9 (o clic derecho → Actualizar "
              "campos) para generar el índice.")
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    doc.add_page_break()


# ── secciones de prueba (catálogo YAML) ─────────────────────────────────────
_SUBTITULOS = ("Objetivo", "Condiciones y desarrollo de la prueba",
               "Señales requeridas", "Criterio de aceptación", "Resultados")


def _seccion_prueba(doc: Document, entrada: dict, cfg: dict,
                    proyecto: ProyectoProtocolo) -> None:
    numero = _numero_de(entrada["seccion"], cfg["offset_numero"])
    aplica = numero in proyecto.pruebas_aplican
    reemplazos = proyecto._reemplazos()
    titulo = entrada["titulo"].strip()
    num_seccion = f"{cfg['capitulo']}.{entrada['seccion']}"

    if not aplica:
        _h2(doc, num_seccion, f"{titulo} — NO APLICA", muted=True)
        nota = proyecto.notas.get(numero, "").strip()
        if nota:  # solo texto entregado por el usuario, nunca redactado aquí
            p = doc.add_paragraph()
            _texto(p, nota, size=9.5, color=MUTED)
        return

    _h2(doc, num_seccion, titulo)
    for item in entrada.get("flujo", []):
        if item.get("tabla") == "aplicabilidad":
            _tabla_aplicabilidad_v2(doc, proyecto.tipo)
        elif "placeholder" in item:
            _figura_placeholder(
                doc, _aplicar_reemplazos(item["placeholder"], reemplazos))
        elif "tabla_datos" in item:
            _tabla_v2(doc, item["tabla_datos"])
        elif "estilo" in item:
            texto = _aplicar_reemplazos(item.get("texto", ""), reemplazos)
            if not texto.strip():
                continue
            p = doc.add_paragraph()
            if (texto.strip() in _SUBTITULOS
                    or item["estilo"] in ("Heading 3", "Subtitle")):
                p.style = doc.styles["Heading 3"]
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                _texto(p, texto, font=DISPLAY, size=10.5, bold=True,
                       color=ELECTRIC_DARK)
            else:
                p.paragraph_format.space_after = Pt(3)
                _texto(p, texto, size=10, color=BODY)


def _pie_pagina(doc: Document, proyecto: ProyectoProtocolo) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    etiqueta = " · ".join(x for x in (
        proyecto.proyecto or proyecto.nombre_central,
        "Protocolo de Pruebas", "Código de Red 2.0") if x)
    _texto(p, f"{etiqueta} · pág. ", size=8, color=LABEL)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


# ── entrada principal ───────────────────────────────────────────────────────
def generar_protocolo_v2(clase: str, proyecto: ProyectoProtocolo,
                         destino: Path) -> Path:
    """Genera el protocolo v2 completo: estructura íntegra de la plantilla
    del usuario (nada se elimina) con el tratamiento visual aprobado.

    `clase`: 'central' (pruebas 21–45) o 'unidad' (pruebas 1–20).
    `destino`: carpeta o ruta .docx.
    """
    if clase not in _CFG:
        raise ValueError(f"clase debe ser 'central' o 'unidad', no {clase!r}")
    cfg = _CFG[clase]
    catalogo = _cargar_catalogo(clase)
    items = _extraer_plantilla(clase)
    reemplazos = proyecto._reemplazos()

    doc = Document()
    _estilos_encabezados(doc)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

    # zona previa al primer H1: portada + control de revisiones
    idx_primer_h1 = next(i for i, (t, _) in enumerate(items) if t == "h1")
    previos = items[:idx_primer_h1]
    alcance_pill = next((filas[0][0].splitlines()[0].strip()
                         for t, filas in previos
                         if t == "tabla" and len(filas) == 1
                         and len(filas[0]) == 1 and filas[0][0].strip()), "")
    _portada(doc, proyecto, previos, alcance_pill)

    tabla_control = next((filas for t, filas in previos
                          if t == "tabla" and filas
                          and filas[0][0].strip().upper() == "REV"), None)
    if tabla_control:
        _h1(doc, "CONTROL DE REVISIONES")
        _tabla_v2(doc, tabla_control)
    _tabla_contenido(doc)

    # capítulos en el orden original de la plantilla
    en_pruebas = False       # dentro del capítulo de pruebas (usa el catálogo)
    tras_h2_pruebas = False  # ya pasó la primera sección vieja → se omite
    for tipo_item, contenido in items[idx_primer_h1:]:
        if tipo_item == "h1":
            if contenido == cfg["inicio_h1"]:
                en_pruebas, tras_h2_pruebas = True, False
                _h1(doc, contenido)
                continue
            if en_pruebas:
                # fin del capítulo de pruebas: insertar secciones del catálogo
                for entrada in catalogo["pruebas"]:
                    if "grupo" in entrada:
                        _h2(doc, "", entrada["grupo"])
                    else:
                        _seccion_prueba(doc, entrada, cfg, proyecto)
                en_pruebas = False
            _h1(doc, contenido)
            continue
        if en_pruebas:
            if tipo_item == "h2":
                tras_h2_pruebas = True
            if tras_h2_pruebas:
                continue  # el contenido viejo lo sustituye el catálogo YAML
            if tipo_item == "tabla":
                if _es_tabla_maestra(contenido):
                    _tabla_maestra_v2(doc, proyecto, clase)
                elif not _es_tabla_aplicabilidad(contenido):
                    _tabla_v2(doc, contenido)
                continue
        # contenido normal (preámbulo y capítulos finales)
        if tipo_item == "h2":
            _h2(doc, "", _aplicar_reemplazos(contenido, reemplazos))
        elif tipo_item == "p":
            _parrafo(doc, _aplicar_reemplazos(contenido, reemplazos))
        elif tipo_item == "bullet":
            _parrafo(doc, _aplicar_reemplazos(contenido, reemplazos),
                     bullet=True)
        elif tipo_item == "figura":
            _figura_placeholder(doc, _aplicar_reemplazos(contenido, reemplazos))
        elif tipo_item == "tabla":
            if _es_tabla_maestra(contenido):
                _tabla_maestra_v2(doc, proyecto, clase)
            elif not _es_tabla_aplicabilidad(contenido):
                _tabla_v2(doc, contenido)

    # si el capítulo de pruebas era el último, insertar el catálogo al final
    if en_pruebas:
        for entrada in catalogo["pruebas"]:
            if "grupo" in entrada:
                _h2(doc, "", entrada["grupo"])
            else:
                _seccion_prueba(doc, entrada, cfg, proyecto)

    _pie_pagina(doc, proyecto)

    destino = Path(destino)
    if destino.suffix.lower() != ".docx":
        destino.mkdir(parents=True, exist_ok=True)
        destino = (destino /
                   f"Protocolo_{clase.capitalize()}_v2_{proyecto.codigo or 'GCV'}.docx")
    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    return destino
