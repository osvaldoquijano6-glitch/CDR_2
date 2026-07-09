"""Plan de Trabajo CRE para Centros de Carga (Cap. 4.1, Manual CONE).

Nota de reestructuración v3: el resto del paquete documental (checklist,
protocolos y anexo de revisiones) se genera ahora en `gcv.protocolos` por
cirugía directa sobre las plantillas del usuario, garantizando formato
idéntico. Este módulo conserva únicamente el Plan de Trabajo de Centros de
Carga, que no tiene plantilla del usuario todavía.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from gcv.models import Installation


def generar_plan_trabajo(inst: Installation, destino: Path) -> Path:
    """Plan de Trabajo CRE (formato del Capítulo 4.1, Manual CONE).

    Documento a presentar ante la CRE. Reproduce la estructura oficial con
    campos en blanco para captura; prellenados los datos que la instalación ya
    aporta (nivel de tensión, demanda contratada) y las tablas de
    requerimientos técnicos con los rangos obligatorios. Sin nombres de
    empresa: placeholders. `destino` puede ser carpeta o ruta .docx.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    destino = Path(destino)
    if destino.suffix.lower() != ".docx":
        destino.mkdir(parents=True, exist_ok=True)
        destino = destino / "Plan_de_Trabajo_CRE.docx"
    destino.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLAN DE TRABAJO\nCÓDIGO DE RED — CONEXIÓN DE CENTROS DE CARGA\n")
    run.bold = True; run.font.size = Pt(20); run.font.color.rgb = RGBColor(0x2A, 0x3F, 0x54)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"{inst.nombre}\nFormato del Capítulo 4.1 — Manual Regulatorio de "
               f"Conexión de Centros de Carga\n{datetime.now():%d/%b/%Y}").font.size = Pt(12)
    doc.add_page_break()

    def campo(etiqueta: str, valor: str = "") -> None:
        par = doc.add_paragraph()
        par.add_run(f"{etiqueta}: ").bold = True
        par.add_run(valor or "________________________________")

    def tabla(headers: list[str], filas: list[list[str]]) -> None:
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h; c.paragraphs[0].runs[0].bold = True
        for fila in filas:
            cells = t.add_row().cells
            for i, v in enumerate(fila):
                cells[i].text = str(v)

    v_kv = inst.tension_poi_kv
    nivel = ("Alta Tensión" if (v_kv or 0) >= 69 else "Media Tensión") if v_kv else ""

    doc.add_heading("1. Datos de identificación del Centro de Carga", level=1)
    campo("1.1 Nivel de tensión en el Punto de Conexión", nivel)
    campo("1.2 Registro de Usuario (RMU/RPU)")
    campo("1.3 Persona física / moral")
    campo("1.3.1 Nombre o razón social del Centro de Carga", "{{NOMBRE_CC}}")
    campo("1.4 Demanda contratada (kW)")
    campo("1.6 Ubicación (domicilio, C.P., municipio, entidad)")
    campo("1.7 Actividad industrial — Código SCIAN")

    doc.add_heading("2. Acreditación de la Representación Legal", level=1)
    campo("2.1 Representante legal (nombre y apellidos)")
    campo("2.2 Domicilio para oír y recibir notificaciones")
    doc.add_paragraph("Nota: anexar instrumento público que acredite la personalidad "
                      "y facultades del representante legal, e identificación oficial.")

    doc.add_heading("3. Requerimientos técnicos del Código de Red aplicables", level=1)
    doc.add_paragraph("3.1 Tensión — rangos obligatorios (Tablas 2.1.A/B):").runs[0].bold = True
    tabla(["Condición", "Rango", "Tiempo"],
          [["Permanente", "95 % – 105 % de Vnom", "Continuo"],
           ["Temporal", "90 % – 110 % de Vnom", "Hasta 20 minutos"]])
    campo("Tensión nominal del Punto de Conexión (kV)",
          f"{v_kv:g}" if v_kv else "")
    doc.add_paragraph("3.2 Frecuencia — rangos obligatorios (Tabla 2.2.A):").runs[0].bold = True
    tabla(["Tiempo", "Frecuencia mínima", "Frecuencia máxima"],
          [["Permanente", "59.0 Hz", "61.0 Hz"],
           ["30 minutos", "58.0 Hz", "62.5 Hz"]])
    doc.add_paragraph("Maniobras de conexión/desconexión: no deben causar desviaciones "
                      "de frecuencia mayores a 0.1 Hz.")
    campo("3.3 Corto circuito — Icc trifásica/monofásica informada por CENACE/Distribuidor")
    doc.add_paragraph("3.4 Factor de potencia — requerimiento vigente:").runs[0].bold = True
    doc.add_paragraph("0.95 en atraso a 1.0 (hasta el 8 de abril de 2026); 0.97 en atraso "
                      "a 1.0 a partir de esa fecha. Medición cinco-minutal, cumplimiento "
                      "≥ 95 % del tiempo en periodo mensual (NOM-001-CRE/SCFI-2019).")
    tabla(["Indicador", "Mínimo", "Percentil 5", "Promedio", "Máximo"],
          [["Factor de potencia", "", "", "", ""]])
    doc.add_paragraph("3.7 Calidad de la potencia — límites obligatorios:").runs[0].bold = True
    tabla(["Indicador", "Límite", "Criterio estadístico"],
          [["Desbalance de tensión", "2 %", "P95 semanal, agregación 10 min"],
           ["Desbalance de corriente", "15 %", "Promedio semanal, agregación 10 min"],
           ["TDD de corriente", "según Tablas 2.8.A/B/C (por Icc/IL y nivel de tensión)",
            "P95 semanal"],
           ["Flicker Pst", "1.0", "P95 semanal"],
           ["Flicker Plt", "0.8", "P95 semanal"]])

    doc.add_heading("4. Plan de Trabajo", level=1)
    doc.add_paragraph(
        "Análisis y estrategia prevista para asegurar el cumplimiento del Código de "
        "Red. Debe incluir: acciones a implementar y análisis de alternativas (equipos "
        "evaluados, retos técnicos y económicos). La futura realización de estudios de "
        "diagnóstico no se considera una acción del Plan de Trabajo.")
    for _ in range(3):
        doc.add_paragraph("________________________________________________________________")

    doc.add_heading("4.1 Cronograma", level=1)
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Acciones previstas", "Fecha de inicio", "Fecha de terminación"]):
        c = t.rows[0].cells[i]; c.text = h; c.paragraphs[0].runs[0].bold = True
    for _ in range(8):
        t.add_row()
    doc.add_paragraph("\n\n____________________________________________")
    doc.add_paragraph("Nombre y firma del Representante Legal para efectos del Código de Red")

    doc.save(str(destino))
    return destino
