"""Informe Word (python-docx): mismas secciones que el HTML, gráficas como PNG
vía kaleido cuando está disponible (si no, se referencia el informe HTML)."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from gcv.reporting.context import ReportContext

_STATUS_COLOR = {
    "CUMPLE": RGBColor(0x0A, 0x5C, 0x0A),
    "NO_CUMPLE": RGBColor(0x8F, 0x1F, 0x1F),
    "NO_EVALUABLE": RGBColor(0x7A, 0x5A, 0x00),
    "PENDIENTE_DOCUMENTAL": RGBColor(0x44, 0x44, 0x44),
}


def _table(doc: Document, headers: list[str], rows: list[list]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)


def _fig_png(fig) -> bytes | None:
    try:
        return fig.to_image(format="png", width=1000, height=520, scale=2)
    except Exception:
        return None  # kaleido/chrome no disponible


def export_docx(ctx: ReportContext, path: Path) -> Path:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    doc.add_heading("Informe técnico de verificación — Código de Red", level=0)
    inst = ctx.installation
    _table(doc, ["Campo", "Valor"], [
        ["Proyecto", ctx.proyecto],
        ["Instalación", f"{inst.nombre} ({inst.kind.value})"],
        ["Clasificación", f"Categoría {inst.category.value if inst.category else 'pendiente'} · "
                          f"{inst.tech.value if inst.tech else '—'}"],
        ["Fecha de emisión", ctx.fecha.strftime("%Y-%m-%d %H:%M")],
        ["Responsable", ctx.responsable or "—"],
    ])

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(ctx.objetivo)
    if ctx.alcance:
        doc.add_heading("2. Alcance", level=1)
        doc.add_paragraph(ctx.alcance)
    doc.add_heading("3. Metodología", level=1)
    doc.add_paragraph(ctx.metodologia)

    doc.add_heading("4. Archivos analizados", level=1)
    _table(doc, ["Fuente", "SHA-256", "Filas", "fs detectada"], [
        [ds.source_path, (ds.source_sha256 or "")[:16] + "…", ds.quality.n_filas,
         f"{ds.quality.fs_detectada_hz:.4g} Hz" if ds.quality.fs_detectada_hz else "—"]
        for ds in ctx.datasets])

    doc.add_heading("5. Resumen de resultados", level=1)
    _table(doc, list(ctx.resumen.keys()), [list(ctx.resumen.values())])

    doc.add_heading("6. Resultados por prueba", level=1)
    sin_png = False
    for r in ctx.resultados:
        doc.add_heading(f"{r.test_id} — {r.test_name}", level=2)
        p = doc.add_paragraph()
        run = p.add_run(f"Resultado: {r.status.value.replace('_', ' ')}")
        run.bold = True
        run.font.color.rgb = _STATUS_COLOR.get(r.status.value, RGBColor(0, 0, 0))
        refs = "; ".join(f"{n.documento} {n.numeral or '(numeral pendiente)'}"
                         for n in r.normative_reference) or "—"
        doc.add_paragraph(f"Referencia normativa: {refs} · Estado del criterio: "
                          f"{r.estado_normativo}")

        if r.pass_fail_details:
            _table(doc, ["Criterio", "Medido", "Límite", "Unidad", "Cumple", "Detalle"], [
                [c.nombre, c.valor_medido,
                 f"{c.comparacion or ''} {c.limite if c.limite is not None else '—'}",
                 c.unidad, {True: "SÍ", False: "NO"}.get(c.cumple, "NO EVALUABLE"),
                 c.detalle] for c in r.pass_fail_details])
        if r.measured_values:
            _table(doc, ["Variable medida", "Valor", "Unidad", "Detalle"], [
                [m.nombre, None if m.valor is None else f"{m.valor:.6g}", m.unidad, m.detalle]
                for m in r.measured_values])
        for w in r.warnings:
            doc.add_paragraph(f"⚠ {w}")
        for fig in ctx.figuras.get(r.test_id, []):
            png = _fig_png(fig)
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(6.5))
            else:
                sin_png = True
        doc.add_paragraph(f"Conclusión: {r.conclusion}")

    doc.add_heading("7. Pendientes", level=1)
    if ctx.pendientes:
        _table(doc, ["ID", "Prueba", "Motivo"], [
            [r.test_id, r.test_name, "; ".join(r.warnings) or r.conclusion]
            for r in ctx.pendientes])
    else:
        doc.add_paragraph("Sin pendientes.")

    doc.add_heading("8. Trazabilidad", level=1)
    for ds in ctx.datasets:
        doc.add_paragraph(f"{ds.source_path} — SHA-256 {(ds.source_sha256 or '')[:16]}…")
        _table(doc, ["Columna original", "Señal", "Unidad orig.", "Factor", "Método"], [
            [m.columna_original, m.senal_canonica, m.unidad_original,
             m.factor_conversion, m.metodo.value] for m in ds.mappings])
    if sin_png:
        doc.add_paragraph(
            "Nota: las gráficas interactivas no pudieron rasterizarse en este "
            "entorno; consulte el informe HTML adjunto.")

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
