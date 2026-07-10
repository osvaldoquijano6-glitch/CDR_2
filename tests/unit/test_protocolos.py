"""Tests del generador de protocolos fiel a las plantillas del usuario."""

import openpyxl
import pytest
from docx import Document

from gcv.protocolos import (ProyectoProtocolo, generar_checklist,
                            generar_protocolo, generar_protocolo_v2,
                            generar_revisiones, universo_pruebas)


@pytest.fixture
def proyecto():
    return ProyectoProtocolo(
        nombre_central="Central Eléctrica Demo Norte",
        codigo="DMN", proyecto="GCE999_DMN_Demo", tipo="C",
        tecnologia="ASINCRONA", ubicacion="Hermosillo, Sonora, México",
        fecha_pruebas="Agosto 2026", fecha_envio="Septiembre 2026",
        pruebas_aplican={1, 2, 3, 9, 17, 21, 22, 23, 26, 30, 45},
        notas={27: "No se cuenta con la infraestructura"},
    )


def test_universo_es_el_del_checklist_usuario():
    uni = universo_pruebas()
    assert len(uni) == 45
    assert uni[0]["numero"] == 1 and "Tensión" in uni[0]["nombre"]
    assert uni[-1]["numero"] == 45
    assert {p["tipo"] for p in uni} <= {"EN SITIO", "DOCUMENTAL", "NA", ""}


def test_protocolo_central_fiel(proyecto, tmp_path):
    ruta = generar_protocolo("central", proyecto, tmp_path)
    doc = Document(str(ruta))

    # portada estampada con los datos del proyecto
    portada = [p.text.strip() for p in doc.paragraphs[:40] if p.text.strip()]
    assert "PROTOCOLO DE PRUEBAS PARA PUESTA EN SERVICIO DE CENTRALES ELÉCTRICAS ASÍNCRONAS" in portada
    assert "Central Tipo C" in portada
    assert any("GCE999_DMN_Demo" in t for t in portada)

    # encabezado con el código del proyecto (no el de la plantilla)
    assert "DMN" in doc.sections[0].header.paragraphs[0].text
    assert "CNK" not in doc.sections[0].header.paragraphs[0].text

    # capítulo de pruebas: 18 secciones; la 3.1 aplica (con contenido) y la
    # 3.7 no aplica (stub con nota del proyecto)
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    numeradas = [h for h in h2 if h.startswith("3.")]
    assert len(numeradas) == 18
    assert numeradas[0] == "3.1. Prueba 1 - Razón de cambio 2.5 Hz/s"
    assert "(no aplica)" in numeradas[6]
    textos = "\n".join(p.text for p in doc.paragraphs)
    assert "No aplica para la Central Eléctrica Demo Norte. No se cuenta con la infraestructura." in textos
    # texto íntegro de una prueba aplicable
    assert "RoCoF = Δf / Δt = 0.5 Hz / 0.200 s = 2.5 Hz/s" in textos

    # tabla maestra con APLICA actualizado (21 SI, 28 No Aplica, 27 nota)
    maestra = next(t for t in doc.tables if "APLICA (SI/NO)" in t.rows[0].cells[-1].text)
    por_numero = {}
    for fila in maestra.rows[2:]:
        num = fila.cells[0].text.strip()
        if num.isdigit():
            por_numero[int(num)] = fila.cells[-1].text.strip()
    assert por_numero[21] == "SI"
    assert por_numero[28] == "No Aplica"
    assert por_numero[27] == "No se cuenta con la infraestructura"


def test_protocolo_unidad_grupos_y_numeracion(proyecto, tmp_path):
    ruta = generar_protocolo("unidad", proyecto, tmp_path)
    doc = Document(str(ruta))
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    # los 4 encabezados de grupo del formato original se conservan
    grupos = [h for h in h2 if not h.startswith("4.")]
    assert len(grupos) == 4
    assert "Sistema de Control de Tensión, Pruebas en Vacio" in grupos
    # numeración sin doble punto y subsecciones anidadas intactas
    assert "4.1. Rango Operativo de Tensión" in h2
    assert "4.2.1. Escalón de Tensión 5%" in h2
    assert not any(".." in h for h in h2)
    # prueba 2 aplica → 4.2.x sin "(no aplica)"; prueba 6 (PSS) no aplica
    assert all("no aplica" not in h.lower() for h in h2 if h.startswith("4.2."))
    assert any(h.startswith("4.6.") and "no aplica" in h.lower() for h in h2)


def test_protocolo_marca_tipo_en_aplicabilidad(proyecto, tmp_path):
    doc = Document(str(generar_protocolo("central", proyecto, tmp_path)))
    aplicables = [t for t in doc.tables if "LA PRUEBA APLICA" in t.rows[0].cells[0].text]
    assert aplicables, "deben conservarse las tablas de aplicabilidad"
    marcas = [c.text.strip() for c in aplicables[0].rows[2].cells]
    assert marcas[4] == "X" and set(marcas) == {"", "X"}  # tipo C


def test_checklist_formato_intacto(tmp_path):
    ruta = generar_checklist(tmp_path / "chk.xlsx",
                             aplica={5: "No Aplica"},
                             comentarios_rev1={5: "Se revisará en sitio"})
    ws = openpyxl.load_workbook(ruta)["DATOS POR PRUEBA"]
    # formato del usuario intacto: encabezado índigo Montserrat y merges
    assert ws["A1"].fill.start_color.rgb == "FF312E81"
    assert ws["A1"].font.name == "Montserrat" and ws["A1"].font.bold
    assert len(ws.merged_cells.ranges) == 12
    assert ws.row_dimensions[1].height == 83.5
    # solo cambian las celdas pedidas; el resto conserva el valor original
    fila5 = next(r for r in ws.iter_rows(min_row=2) if r[3].value == 5)
    assert fila5[7].value == "No Aplica"
    assert fila5[9].value == "Se revisará en sitio"
    fila1 = next(r for r in ws.iter_rows(min_row=2) if r[3].value == 1)
    assert fila1[7].value == "SI"


def test_revisiones_formato_usuario(tmp_path):
    ruta = generar_revisiones(tmp_path / "rev.xlsx", proyecto="Central Demo Norte")
    wb = openpyxl.load_workbook(ruta)
    assert wb.sheetnames == ["Portada", "Comentarios CENACE"]
    textos = [str(c.value) for fila in wb["Portada"].iter_rows() for c in fila if c.value]
    assert any(t == "Proyecto: Central Demo Norte" for t in textos)
    assert wb["Comentarios CENACE"]["A4"].value == "No."


def test_clase_invalida(proyecto, tmp_path):
    with pytest.raises(ValueError):
        generar_protocolo("subestacion", proyecto, tmp_path)


def _tabla_maestra_v2(doc):
    return next(t for t in doc.tables
                if len(t.columns) == 4 and t.rows[0].cells[0].text.strip() == "No.")


def test_protocolo_v2_central(proyecto, tmp_path):
    ruta = generar_protocolo_v2("central", proyecto, tmp_path)
    doc = Document(str(ruta))
    titulo = next(p.text for p in doc.paragraphs if "PROTOCOLO" in p.text)
    assert "ASÍNCRONAS" in titulo  # tecnología del proyecto en el título
    maestra = _tabla_maestra_v2(doc)
    assert len(maestra.rows) - 1 == 25  # universo por central: 21–45
    numeros = [int(r.cells[0].text) for r in maestra.rows[1:]]
    assert numeros == list(range(21, 46))
    estados = {r.cells[3].text.strip() for r in maestra.rows[1:]
               if int(r.cells[0].text) != 27}
    assert estados <= {"SI", "NO APLICA"}  # regla dura: sin notas inventadas
    fila27 = next(r for r in maestra.rows[1:] if r.cells[0].text == "27")
    assert fila27.cells[3].text.strip() == "No se cuenta con la infraestructura"


def test_protocolo_v2_conserva_estructura_completa(proyecto, tmp_path):
    """Reestructurar no es eliminar: todos los capítulos H1 de la plantilla
    original deben existir en el v2, y las fuentes deben ser universales."""
    import re
    import zipfile

    from gcv.config.settings import NORMATIVE_DIR

    def norm(s):
        return re.sub(r"\W", "", s).lower()

    for clase in ("central", "unidad"):
        ruta = generar_protocolo_v2(clase, proyecto, tmp_path / clase)
        plantilla = (NORMATIVE_DIR / "plantillas_usuario" /
                     f"Protocolo_{clase.capitalize()}_v1.2.1.docx")
        h1_orig = [p.text.strip() for p in Document(str(plantilla)).paragraphs
                   if p.style.name == "Heading 1"]
        gen = Document(str(ruta))
        h1_gen = [norm(p.text) for p in gen.paragraphs
                  if p.style.name == "Heading 1"]
        faltantes = [h for h in h1_orig if norm(h) not in h1_gen]
        assert not faltantes, f"{clase}: capítulos eliminados {faltantes}"
        # tablas de definiciones y abreviaciones conservadas
        assert any("Término" in t.rows[0].cells[0].text for t in gen.tables)
        assert any("Abreviación" in t.rows[0].cells[0].text for t in gen.tables)
        # fuentes universales: sin tipografías que el revisor deba instalar
        xml = zipfile.ZipFile(str(ruta)).read("word/document.xml").decode()
        assert "Montserrat" not in xml and "Helvetica" not in xml
        assert 'w:ascii="Arial"' in xml and 'w:ascii="Calibri"' in xml


def test_protocolo_v2_unidad_sin_notas(tmp_path):
    p = ProyectoProtocolo(nombre_central="Demo", codigo="DMO", tipo="B",
                          tecnologia="SINCRONA", pruebas_aplican={1, 2, 6})
    ruta = generar_protocolo_v2("unidad", p, tmp_path)
    maestra = _tabla_maestra_v2(Document(str(ruta)))
    assert len(maestra.rows) - 1 == 20  # universo por unidad: 1–20
    estados = [r.cells[3].text.strip() for r in maestra.rows[1:]]
    assert set(estados) == {"SI", "NO APLICA"}
    assert estados.count("SI") == 3
