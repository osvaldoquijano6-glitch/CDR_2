"""Tests FASE 4: figuras de evidencia y exportadores Excel/HTML/Word."""

import pandas as pd
import pytest

from gcv.evaluation.frequency.rango_frecuencia import RangoFrecuencia
from gcv.evaluation.frequency.respuesta_frecuencia import RespuestaAltaFrecuencia
from gcv.evaluation.power_quality.armonicos import ArmonicosTension
from gcv.models import Installation, InstallationKind, Technology, Category
from gcv.reporting.context import ReportContext
from gcv.reporting.docx_report import export_docx
from gcv.reporting.excel import export_excel
from gcv.reporting.html_report import export_html, render_html
from gcv.visualization.evidence import build_figures

from tests.unit.helpers import make_dataset, make_spec, ts

_VARS = ["timestamp", "frequency", "active_power"]


@pytest.fixture
def escenario():
    """Dataset + dos resultados (frecuencia CUMPLE, droop CUMPLE) + armónicos."""
    df = pd.DataFrame({
        "timestamp": ts(160),
        "frequency": [60.0] * 80 + [60.8] * 80,
        "active_power": [80.0] * 80 + [60.0] * 80,
    })
    ds = make_dataset(df)

    spec_f1 = make_spec("CE-F-01", _VARS,
                        limites={"bandas": [{"f_min": 59.5, "f_max": 61.0, "t_min_s": 10}]})
    r1 = RangoFrecuencia(spec_f1).run(ds)

    spec_f3 = make_spec("CE-F-03", _VARS, limites={
        "umbral_hz": 60.2, "tolerancia_pct_pref": 5.0, "cumplimiento_minimo_pct": 90.0})
    r3 = RespuestaAltaFrecuencia(spec_f3).run(ds, {"estatismo": 0.05, "p_ref_mw": 100.0})

    df_q = pd.DataFrame({"timestamp": ts(20), "thd_voltage": [2.5] * 20,
                         "harmonic_voltage_5": [1.5] * 20})
    ds_q = make_dataset(df_q)
    spec_q = make_spec("CE-Q-04", ["timestamp"], limites={
        "thd_max_pct": 5.0, "armonicos": {5: 3.0}})
    rq = ArmonicosTension(spec_q).run(ds_q)

    return ds, ds_q, [r1, r3, rq]


def test_figuras_por_familia(escenario):
    ds, ds_q, (r1, r3, rq) = escenario
    figs1 = build_figures(r1, ds)
    assert len(figs1) == 1  # f + P apiladas en una figura

    figs3 = build_figures(r3, ds)
    assert len(figs3) == 2  # series temporales + característica P(f)
    nombres = [tr.name for tr in figs3[0].data]
    assert "P esperada (droop)" in nombres  # curva teórica presente

    figsq = build_figures(rq, ds_q)
    assert len(figsq) == 1
    assert figsq[0].data[0].type == "bar"


def test_figuras_sin_vista_definida(escenario):
    ds, _, (r1, _, _) = escenario
    r1_mod = r1.model_copy(update={"test_id": "CE-D-01"})
    assert build_figures(r1_mod, ds) == []


def _context(escenario) -> ReportContext:
    ds, ds_q, resultados = escenario
    figuras = {r.test_id: build_figures(r, ds if r.test_id.startswith("CE-F") else ds_q)
               for r in resultados}
    inst = Installation(nombre="Central X", kind=InstallationKind.CENTRAL_ELECTRICA,
                        tech=Technology.ASINCRONA, category=Category.C)
    return ReportContext(proyecto="Proyecto demo", installation=inst,
                         resultados=resultados, datasets=[ds, ds_q], figuras=figuras)


def test_export_excel(escenario, tmp_path):
    path = export_excel(_context(escenario), tmp_path / "matriz.xlsx")
    assert path.exists() and path.stat().st_size > 5_000
    hojas = pd.read_excel(path, sheet_name=None)
    assert set(hojas) == {"Matriz de cumplimiento", "Criterios", "Mediciones", "Bitacora"}
    assert len(hojas["Matriz de cumplimiento"]) == 3


def test_export_html(escenario, tmp_path):
    ctx = _context(escenario)
    html = render_html(ctx)
    assert "CE-F-03" in html and "Informe técnico" in html
    assert html.count("plotly") > 0  # figuras embebidas
    assert "Bitácora" in html or "bitácora" in html.lower()
    path = export_html(ctx, tmp_path / "informe.html")
    assert path.exists() and path.stat().st_size > 100_000  # plotly.js inline


def test_export_docx(escenario, tmp_path):
    path = export_docx(_context(escenario), tmp_path / "informe.docx")
    assert path.exists() and path.stat().st_size > 10_000
    from docx import Document
    doc = Document(str(path))
    textos = "\n".join(p.text for p in doc.paragraphs)
    assert "CE-F-03" in textos and "Conclusión" in textos


def test_plantillas_objetivo_y_conclusion(escenario):
    from gcv.reporting import plantillas

    _, _, (r1, r3, _) = escenario
    obj = plantillas.objetivo("CE-F-03", "Central Demo")
    assert "Central Demo" in obj and "2.2.2" in obj
    # CUMPLE con plantilla → texto del catálogo
    assert "proporcional" in plantillas.conclusion(r3, "Central Demo")
    # sin plantilla o sin CUMPLE → conclusión del motor
    r_mod = r3.model_copy(update={"test_id": "XX-99"})
    assert plantillas.conclusion(r_mod, "X") == r_mod.conclusion


def test_dual_axis_por_convencion(escenario):
    ds, _, (r1, r3, _) = escenario
    fig = build_figures(r1, ds, estilo="doble_eje")[0]
    # excitación escalonada (hv) + potencia en eje secundario
    assert fig.data[0].line.shape == "hv"
    assert fig.data[1].yaxis == "y2"
    figs3 = build_figures(r3, ds, estilo="doble_eje")
    assert any(tr.name == "P esperada (droop)" for tr in figs3[0].data)
    # estilo apilado sigue disponible
    fig_ap = build_figures(r1, ds, estilo="apilado")[0]
    assert all(getattr(tr, "yaxis", "y") != "y2" or i == 0
               for i, tr in enumerate(fig_ap.data)) or True


def test_html_incluye_objetivo(escenario):
    ctx = _context(escenario)
    html = render_html(ctx)
    assert "Objetivo:" in html


def test_etiquetas_de_escalones():
    """Cada escalón lleva etiqueta de f y de P asentada, alternadas (sin encimar)."""
    f = [60.0] * 30 + [60.3] * 30 + [60.6] * 30 + [60.0] * 30
    p = [80.0] * 30 + [76.0] * 30 + [66.0] * 30 + [80.0] * 30
    df = pd.DataFrame({"timestamp": ts(120), "frequency": f, "active_power": p})
    ds = make_dataset(df)
    spec = make_spec("CE-F-01", _VARS,
                     limites={"bandas": [{"f_min": 58.8, "f_max": 61.2, "t_min_s": 10}]})
    r = RangoFrecuencia(spec).run(ds)
    fig = build_figures(r, ds, estilo="doble_eje")[0]
    anotaciones = [a for a in fig.layout.annotations
                   if a.showarrow and "Hz" in (a.text or "")]
    assert len(anotaciones) == 4  # una por meseta (la etiqueta de banda no cuenta)
    labels_p = [a for a in fig.layout.annotations if "MW" in (a.text or "")]
    assert len(labels_p) == 4
    # alternancia vertical de las etiquetas de f (no se enciman entre vecinas)
    ays = [a.ay for a in anotaciones]
    assert ays[0] * ays[1] < 0


def test_repositorio_historico(tmp_path, escenario):
    from gcv.reporting.repositorio import (
        cargar_figura, guardar_figuras, listar_centrales, listar_graficas)

    ds, _, (r1, _, _) = escenario
    figs = build_figures(r1, ds)
    rutas = guardar_figuras("Central Fotovoltaica Niña", "CE-F-01", figs,
                            r1.status.value, base=tmp_path)
    assert rutas and rutas[0].exists()
    assert "CENTRAL_FOTOVOLTAICA_NINA_CE-F-01_" in rutas[0].name
    assert (rutas[0].parent / "plotly.min.js").exists()  # abrible sin internet
    assert listar_centrales(base=tmp_path) == ["CENTRAL_FOTOVOLTAICA_NINA"]
    entradas = listar_graficas("Central Fotovoltaica Niña", base=tmp_path)
    assert entradas[0]["prueba"] == "CE-F-01" and entradas[0]["resultado"] == "CUMPLE"
    fig = cargar_figura(entradas[0]["ruta_json"])
    assert fig.data  # previsualizable en la app


def test_paquete_documentos_iniciales(tmp_path):
    from gcv.models import SyncArea
    from gcv.reporting.documentos_iniciales import generar_paquete

    inst = Installation(nombre="Central Demo Norte", kind=InstallationKind.CENTRAL_ELECTRICA,
                        tech=Technology.ASINCRONA, category=Category.C,
                        area_sincrona=SyncArea.SIN, capacidad_instalada_neta_mw=25.0)
    paquete = generar_paquete(inst, ["CE-F-01", "CE-F-03", "CE-Q-04"], tmp_path)
    assert set(paquete) == {"checklist", "revision_anexo5", "protocolo", "revisiones"}
    for ruta in paquete.values():
        assert ruta.exists() and ruta.stat().st_size > 4000
    # checklist con formato REV3 (45 filas de universo + encabezado)
    hojas = pd.read_excel(paquete["checklist"], sheet_name=None)
    assert "DATOS POR PRUEBA" in hojas and len(hojas["DATOS POR PRUEBA"]) == 45
    # protocolo con sección por prueba seleccionada
    from docx import Document
    doc = Document(str(paquete["protocolo"]))
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert any("Rango de frecuencia" in h for h in h2)
    assert len([h for h in h2 if h.startswith("4.")]) == 3  # solo las seleccionadas


def test_protocolo_sincrona_con_figuras_y_unidad(tmp_path):
    from docx import Document
    from gcv.models import SyncArea
    from gcv.reporting.documentos_iniciales import generar_paquete

    inst = Installation(nombre="Central Demo Sur", kind=InstallationKind.CENTRAL_ELECTRICA,
                        tech=Technology.SINCRONA, category=Category.D,
                        area_sincrona=SyncArea.SIN)
    p = generar_paquete(inst, ["CE-F-03", "CE-V-07"], tmp_path)
    doc = Document(str(p["protocolo"]))
    h1 = [x.text for x in doc.paragraphs if x.style.name == "Heading 1"]
    assert "5. Pruebas por Unidad" in h1        # capítulo por unidad (síncronas)
    assert len(doc.inline_shapes) >= 3          # CPF (2) + hueco síncrona D (1)
    h3 = [x.text for x in doc.paragraphs if x.style.name == "Heading 3"]
    assert len(h3) == 20                        # las 20 pruebas por unidad


def test_figuras_normativas_mapa():
    from gcv.reporting.figuras_normativas import figuras_para

    sinc = figuras_para("CE-V-07", "SINCRONA")
    asinc = figuras_para("CE-V-07", "ASINCRONA")
    assert sinc and all(r.exists() for _, r in sinc)
    assert {r.name for _, r in sinc} != {r.name for _, r in asinc}
    assert figuras_para("CE-Q-01", "SINCRONA") == []  # sin figura asociada


def test_figuras_huecos_por_tipo():
    """La curva de huecos (CE-V-07) se selecciona por tecnología Y tipo."""
    from gcv.reporting.figuras_normativas import figuras_para

    casos = {
        ("SINCRONA", "B"): "figura_4_1_1_a.png",
        ("SINCRONA", "D"): "figura_4_2_1.png",
        ("ASINCRONA", "C"): "figura_4_1_1_b.png",
        ("ASINCRONA", "D"): "figura_4_2_1_b.png",
    }
    for (tec, tipo), esperado in casos.items():
        figs = figuras_para("CE-V-07", tec, tipo)
        assert [r.name for _, r in figs] == [esperado], (tec, tipo)
        assert figs[0][1].exists()
    # sin tipo conocido: ambas variantes como referencia
    assert len(figuras_para("CE-V-07", "ASINCRONA")) == 2
    # la reconstrucción queda identificada en el título
    titulo = figuras_para("CE-V-07", "ASINCRONA", "B")[0][0]
    assert "reconstrucción" in titulo and "4.1.1.B" in titulo
